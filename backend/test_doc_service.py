
import os
import pytest
from services.doc_service import DocService
from docxtpl import DocxTemplate
from docx import Document

def test_doc_service_extract_fields(tmp_path):
    # Create a dummy docx template with tags
    doc_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Hello {{first_name}}, welcome to {{city}}!")
    doc.save(doc_path)
    
    service = DocService()
    fields = service.extract_fields(str(doc_path))
    
    assert len(fields) == 2
    field_names = [f["name"] for f in fields]
    assert "first_name" in field_names
    assert "city" in field_names
    assert fields[0]["type"] == "text"

def test_doc_service_fill_docx(tmp_path):
    # Create a dummy docx template with tags
    input_path = tmp_path / "template.docx"
    output_path = tmp_path / "filled.docx"
    doc = Document()
    doc.add_paragraph("Hello {{name}}!")
    doc.save(input_path)
    
    service = DocService()
    data = {"name": "Antigravity"}
    success = service.fill_docx(str(input_path), data, str(output_path))
    
    assert success is True
    assert os.path.exists(output_path)
    
    # Verify content
    filled_doc = Document(output_path)
    full_text = []
    for para in filled_doc.paragraphs:
        full_text.append(para.text)
    assert "Hello Antigravity!" in " ".join(full_text)
