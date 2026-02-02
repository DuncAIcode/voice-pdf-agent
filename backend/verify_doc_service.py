
import os
import sys

# Add the current directory to sys.path to import services
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from services.doc_service import DocService
from docx import Document

def run_tests():
    print("Starting DocService Verification...")
    service = DocService()
    
    # Test 1: Extraction
    print("Test 1: Extracting fields from docx...")
    doc = Document()
    doc.add_paragraph("Hello {{first_name}}, welcome to {{city}}!")
    test_doc_path = "test_template.docx"
    doc.save(test_doc_path)
    
    try:
        fields = service.extract_fields(test_doc_path)
        print(f"Extracted fields: {[f['name'] for f in fields]}")
        assert "first_name" in [f["name"] for f in fields]
        assert "city" in [f["name"] for f in fields]
        print("Test 1 Passed!")
    except Exception as e:
        print(f"Test 1 Failed: {e}")
        return

    # Test 2: Filling
    print("Test 2: Filling docx template...")
    output_path = "test_filled.docx"
    data = {"first_name": "Antigravity", "city": "DeepMind"}
    
    try:
        success = service.fill_docx(test_doc_path, data, output_path)
        assert success is True
        assert os.path.exists(output_path)
        
        # Verify content
        filled_doc = Document(output_path)
        full_text = " ".join([para.text for para in filled_doc.paragraphs])
        print(f"Filled text: {full_text}")
        assert "Antigravity" in full_text
        assert "DeepMind" in full_text
        print("Test 2 Passed!")
    except Exception as e:
        print(f"Test 2 Failed: {e}")
    finally:
        # Cleanup
        if os.path.exists(test_doc_path): os.remove(test_doc_path)
        if os.path.exists(output_path): os.remove(output_path)

if __name__ == "__main__":
    run_tests()
