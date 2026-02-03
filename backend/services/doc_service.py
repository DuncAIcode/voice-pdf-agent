import os
import json
from typing import List, Dict, Any
from docxtpl import DocxTemplate
from docx import Document
import google.generativeai as genai
import jinja2
import mammoth

class DocService:
    def __init__(self):
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            genai.configure(api_key=api_key)
        else:
            print("Warning: GEMINI_API_KEY not found in DocService.")
        self.model = genai.GenerativeModel('gemini-1.5-flash-8b')

    def extract_fields(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extracts {{tags}} from a .docx template.
        Returns a list of dictionaries with name and label.
        """
        try:
            doc = DocxTemplate(file_path)
            # Find undeclared tags
            tags = doc.get_undeclared_template_variables()
            
            extracted = []
            for tag in tags:
                extracted.append({
                    "name": tag,
                    "label": tag.replace('_', ' ').capitalize(),
                    "type": "text", # Word templates usually expect text
                    "page": 1,      # Page number is harder to determine in Word, defaulting to 1
                    "coordinates": None # No coordinates for Word tags
                })
            
            return extracted
        except Exception as e:
            print(f"Error extracting DOCX fields: {e}")
            return []

    def fill_docx(self, file_path: str, data: Dict[str, Any], output_path: str):
        """
        Fills a .docx template with the provided data using docxtpl.
        """
        try:
            doc = DocxTemplate(file_path)
            doc.render(data)
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error filling DOCX: {e}")
            return False

    def analyze_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extracts text from a regular .docx and uses AI to suggest potential fields.
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text)
            
            # Extract Headers/Footers
            for section in doc.sections:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)

            content = "\n".join(full_text)
            
            if not content.strip():
                print("Warning: No text content found in document for analysis.")
                return []

            prompt = f"""
            Analyze the following document text and identify potential dynamic fields, form placeholders, or areas meant to be filled in.
            
            Look for patterns like:
            1. Explicit placeholders: [Name], <Date>, (Amount), {{Variable}}
            2. Form labels with gaps: "Name: _________", "Date: ____________", "Address: " (followed by space or newline)
            3. Logical entities in context: "This agreement is between [Party A] and [Party B]", "effective as of [Date]"
            4. Capitalized labels that look like field names: "CLIENT NAME:", "PROJECT TOTAL:"
            
            For each field found, provide:
            - "original_text": The EXACT string from the document (including brackets/underscores) to be replaced.
            - "suggested_tag": A generic slug-style variable name (e.g., client_name, project_date).
            - "reason": A brief note on why this was identified.

            DOCUMENT TEXT:
            {content[:4000]}
            
            JSON OUTPUT ONLY. DO NOT INCLUDE MARKDOWN.
            Return an array of objects.
            """
            
            response = self.model.generate_content(prompt)
            text = response.text.strip()
            
            # More robust JSON extraction
            try:
                # Try finding the first [ and last ]
                start = text.find('[')
                end = text.rfind(']') + 1
                
                if start != -1 and end != -1:
                    json_str = text[start:end]
                    suggestions = json.loads(json_str)
                else:
                    # Try as a single object
                    start_obj = text.find('{')
                    end_obj = text.rfind('}') + 1
                    if start_obj != -1 and end_obj != -1:
                        obj = json.loads(text[start_obj:end_obj])
                        suggestions = [obj] if isinstance(obj, dict) else obj
                    else:
                        suggestions = json.loads(text)
                
                print(f"Detected {len(suggestions)} suggestions for {file_path}")
                return suggestions if isinstance(suggestions, list) else []
            except Exception as json_e:
                print(f"JSON Parse Error: {json_e} - Text: {text}")
                return []
        except Exception as e:
            print(f"Error analyzing document with AI: {e}")
            # Fallback to pattern-based analysis
            print("Falling back to pattern-based analysis...")
            return self._analyze_patterns(content)

    def _analyze_patterns(self, text: str) -> List[Dict[str, Any]]:
        """
        Uses regex to find common placeholders: [Name], <Date>, {Var}, and ______
        """
        import re
        suggestions = []
        
        # 1. Square brackets [Field]
        for match in re.finditer(r'\[([^\]]+)\]', text):
            original = match.group(0)
            clean = re.sub(r'[^a-zA-Z0-9]', '_', match.group(1)).lower().strip('_')
            suggestions.append({
                "original_text": original,
                "suggested_tag": clean or "field",
                "reason": "Detected via [brackets] pattern"
            })
            
        # 2. Underscores (minimum 5) ______
        # This is harder because we need context for the tag name
        # We'll look for text followed by underscores: Name: _______
        for match in re.finditer(r'([a-zA-Z\s]{2,20})[:]?\s*_{5,}', text):
            label = match.group(1).strip()
            original = match.group(0)
            # Find just the underscores part of the original
            underscore_match = re.search(r'_{5,}', original)
            if underscore_match:
                just_underscores = underscore_match.group(0)
                clean_label = re.sub(r'[^a-zA-Z0-9]', '_', label).lower().strip('_')
                suggestions.append({
                    "original_text": just_underscores,
                    "suggested_tag": clean_label or "field",
                    "reason": f"Detected underscore placeholder for '{label}'"
                })

        # 3. Curly brackets {{Var}}
        for match in re.finditer(r'\{\{([^\}]+)\}\}', text):
            original = match.group(0)
            clean = re.sub(r'[^a-zA-Z0-9]', '_', match.group(1)).lower().strip('_')
            suggestions.append({
                "original_text": original,
                "suggested_tag": clean or "field",
                "reason": "Detected via {{brackets}} pattern"
            })

        # Deduplicate by original_text
        seen = set()
        unique = []
        for s in suggestions:
            if s["original_text"] not in seen:
                unique.append(s)
                seen.add(s["original_text"])
                
        return unique

    def transform_template(self, input_path: str, output_path: str, replacements: List[Dict[str, str]]):
        """
        Replaces specific strings in a .docx with {{tags}} while preserving styles.
        """
        try:
            doc = Document(input_path)
            
            def process_text_container(container):
                for paragraph in container.paragraphs:
                    for rep in replacements:
                        target = rep["original_text"]
                        tag = f"{{{{{rep['tag_name']}}}}}"
                        
                        if target in paragraph.text:
                            # Try replacing at the run level first to preserve style
                            found_in_run = False
                            for run in paragraph.runs:
                                if target in run.text:
                                    run.text = run.text.replace(target, tag)
                                    found_in_run = True
                            
                            # Fallback: If target is split across runs, we have to do paragraph-level replacement.
                            # This merges formatting, but it ensures the tag is actually injected.
                            if not found_in_run:
                                paragraph.text = paragraph.text.replace(target, tag)
            
            # Process main paragraphs
            process_text_container(doc)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_text_container(cell)

            # Process Headers and Footers
            for section in doc.sections:
                process_text_container(section.header)
                process_text_container(section.footer)
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error transforming template: {e}")
            return False

    def get_document_preview(self, file_path: str) -> str:
        """
        Converts a .docx file to HTML for previewing, preserving basic formatting.
        """
        try:
            with open(file_path, "rb") as docx_file:
                # Use mammoth to convert to HTML
                result = mammoth.convert_to_html(docx_file)
                html = result.value # The generated HTML
                messages = result.messages # Any messages, such as warnings
                
                # Mammoth focuses on structural HTML, which is good for our needs
                # and safer for injecting into the frontend.
                return html
        except Exception as e:
            print(f"Error getting document preview: {e}")
            return f"<p>Error loading preview: {str(e)}</p>"
